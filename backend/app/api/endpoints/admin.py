from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from sqlalchemy.orm import Session
from typing import List, Optional
from datetime import datetime
from pydantic import BaseModel
import uuid
import base64
import logging
import json

from app.core.database import get_db
from app.core.config import settings
from app.core.security import get_current_user, require_admin
from app.models.user import User
from app.models.property import Property
from app.models.order import Order, OrderItem, OrderStatus, OrderItemFlag
from app.models.supplier import Supplier
from app.models.inventory import InventoryItem

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/admin", tags=["Admin"])


# ============== SCHEMAS ==============

class ExtractedOrderItem(BaseModel):
    item_name: str
    quantity: float
    unit: Optional[str] = None
    unit_price: Optional[float] = None
    supplier_name: Optional[str] = None
    category: Optional[str] = None
    notes: Optional[str] = None
    is_recurring: bool = True  # Whether item appears on inventory printout sheets


class ExtractedOrderData(BaseModel):
    items: List[ExtractedOrderItem]
    order_date: Optional[str] = None
    order_number: Optional[str] = None
    total: Optional[float] = None
    supplier_name: Optional[str] = None
    notes: Optional[str] = None
    confidence_score: float = 0.0
    raw_text: Optional[str] = None


class SeedOrderRequest(BaseModel):
    property_id: int
    order_date: str  # ISO format date
    items: List[ExtractedOrderItem]
    status: str = "ordered"  # Default to "ordered" so items can go through receiving workflow
    notes: Optional[str] = None


class SeedOrderResponse(BaseModel):
    id: int
    order_number: str
    property_id: int
    status: str
    week_of: Optional[datetime] = None
    item_count: int
    estimated_total: float
    created_at: datetime

    class Config:
        from_attributes = True


# ============== HELPER FUNCTIONS ==============

def generate_order_number(property_code: str) -> str:
    """Generate order number with property code and date (e.g., YRC-20251215)"""
    return f"{property_code}-{datetime.utcnow().strftime('%Y%m%d')}"


def normalize_item_name(name: str) -> str:
    """
    Normalize an item name for comparison by:
    - Converting to lowercase
    - Removing extra whitespace
    - Handling common variations
    """
    if not name:
        return ""

    # Convert to lowercase and strip
    normalized = name.lower().strip()

    # Remove common filler words and punctuation
    normalized = normalized.replace(",", " ").replace("-", " ").replace("'", "")

    # Collapse multiple spaces
    normalized = " ".join(normalized.split())

    return normalized


def get_name_tokens(name: str) -> set:
    """Get a set of meaningful tokens from a name for fuzzy matching"""
    normalized = normalize_item_name(name)

    # Split into tokens
    tokens = set(normalized.split())

    # Remove very common/short words that don't help with matching
    stop_words = {'the', 'a', 'an', 'of', 'and', 'or', 'for', 'in', 'on', 'lb', 'oz', 'ct', 'pk', 'bag', 'box', 'case'}
    tokens = tokens - stop_words

    return tokens


def calculate_name_similarity(name1: str, name2: str) -> float:
    """
    Calculate similarity between two item names.
    Returns a score from 0.0 to 1.0.

    Handles cases like:
    - "Onions, green" vs "Green Onions"
    - "2% Milk" vs "Milk 2%"
    - "SOCK SALMON" vs "Sockeye Salmon"
    """
    # Exact match (normalized)
    norm1 = normalize_item_name(name1)
    norm2 = normalize_item_name(name2)

    if norm1 == norm2:
        return 1.0

    # Get tokens for each name
    tokens1 = get_name_tokens(name1)
    tokens2 = get_name_tokens(name2)

    if not tokens1 or not tokens2:
        return 0.0

    # Calculate Jaccard similarity (intersection over union)
    intersection = tokens1 & tokens2
    union = tokens1 | tokens2

    jaccard = len(intersection) / len(union) if union else 0.0

    # Also check if one name contains the other (handles abbreviations)
    contains_bonus = 0.0
    if norm1 in norm2 or norm2 in norm1:
        contains_bonus = 0.3

    # Check for common food name inversions (e.g., "Beans, black" vs "Black Beans")
    # If all tokens from one are in the other, it's likely the same item
    if tokens1.issubset(tokens2) or tokens2.issubset(tokens1):
        jaccard = max(jaccard, 0.85)

    return min(1.0, jaccard + contains_bonus)


def find_matching_inventory_item(
    item_name: str,
    property_id: int,
    db,
    similarity_threshold: float = 0.6
) -> tuple:
    """
    Find a matching inventory item by name similarity.

    Returns:
        (inventory_item, similarity_score) if found, (None, 0.0) otherwise
    """
    # Get all inventory items for this property
    inventory_items = db.query(InventoryItem).filter(
        InventoryItem.property_id == property_id,
        InventoryItem.is_active == True
    ).all()

    best_match = None
    best_score = 0.0

    for inv_item in inventory_items:
        score = calculate_name_similarity(item_name, inv_item.name)
        if score > best_score:
            best_score = score
            best_match = inv_item

    if best_score >= similarity_threshold:
        logger.info(f"Matched '{item_name}' to existing inventory item '{best_match.name}' (score: {best_score:.2f})")
        return best_match, best_score

    return None, 0.0


def calculate_order_total(order: Order) -> float:
    """Calculate estimated total for an order"""
    total = 0.0
    for item in order.items:
        qty = item.approved_quantity if item.approved_quantity is not None else item.requested_quantity
        price = item.unit_price or 0
        total += qty * price
    return total


def convert_pdf_to_images(pdf_content: bytes) -> list:
    """Convert PDF pages to base64-encoded PNG images for OpenAI Vision API"""
    try:
        import fitz  # PyMuPDF
        import io

        images = []
        pdf_doc = fitz.open(stream=pdf_content, filetype="pdf")

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            # Render at 2x resolution for better OCR
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)

            # Convert to PNG bytes
            img_bytes = pix.tobytes("png")
            img_base64 = base64.standard_b64encode(img_bytes).decode("utf-8")
            images.append(img_base64)

            # Limit to first 5 pages to avoid token limits
            if page_num >= 4:
                logger.warning(f"PDF has {len(pdf_doc)} pages, only processing first 5")
                break

        pdf_doc.close()
        return images
    except Exception as e:
        logger.error(f"Error converting PDF to images: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to convert PDF to images: {str(e)}"
        )


def extract_text_from_docx(docx_content: bytes) -> str:
    """Extract text content from a DOCX file"""
    try:
        from docx import Document
        import io

        doc = Document(io.BytesIO(docx_content))
        full_text = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))

        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to read DOCX file: {str(e)}"
        )


async def extract_order_with_ai(content: bytes, file_type: str, supplier_names: List[str]) -> ExtractedOrderData:
    """Use OpenAI to extract order data from PDF or DOCX"""
    if not settings.OPENAI_API_KEY:
        raise HTTPException(
            status_code=500,
            detail="OpenAI API key not configured. Please set OPENAI_API_KEY in environment."
        )

    try:
        import openai
        client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)

        # Build supplier list for the prompt
        supplier_list_str = ", ".join(supplier_names) if supplier_names else "No suppliers defined yet"

        system_prompt = f"""You are an expert at extracting purchase order data from documents.
Extract all line items from the order document, including:
- Item name - STANDARDIZE to common/recognizable names
- Quantity ordered
- Unit of measure
- Unit price if visible
- Supplier name - Match to existing suppliers when possible
- Product category

ITEM NAME RULES - VERY IMPORTANT:
- Convert awkward/inverted names to standard common names
- Examples:
  - "Onions, green" → "Green Onions"
  - "Beans, black" → "Black Beans"
  - "Peppers, bell red" → "Red Bell Peppers"
  - "Oikos 000" → "Oikos 000 Yogurt" (add product type if unclear)
  - "Milk 2%" → "2% Milk"
- Keep brand names when present (e.g., "Oikos", "Kirkland")
- Make names clear and recognizable at a glance
- DO NOT add descriptions or notes about what the item is - just name it clearly

EXISTING SUPPLIERS (use exact names when matching):
{supplier_list_str}

SUPPLIER MATCHING RULES:
- If a supplier name in the document looks similar to an existing supplier, use the EXACT existing supplier name
- For example: "COSTCO PRODUCE" or "Costco" should become "Costco Business" if that exists in the list
- "US Foods" variations should match to the US Foods supplier in the list
- Only use null for supplier if you truly cannot determine or match to any existing supplier

PRODUCT CATEGORIES (pick the most appropriate one):
- Produce (fruits, vegetables, fresh herbs)
- Dairy (milk, cheese, butter, eggs, yogurt)
- Protein (meat, poultry, fish, seafood)
- Dry Goods (rice, pasta, flour, canned goods, cereals)
- Beverages (drinks, juices, coffee, tea)
- Frozen (frozen foods, ice cream)
- Bakery (bread, pastries, baked goods)
- Condiments (sauces, dressings, oils, vinegars)
- Spices (spices, seasonings, herbs, salt, pepper)
- Packaged Goods (snacks, chips, crackers, pre-packaged foods)
- Paper Goods (napkins, paper towels, plates)
- Cleaning Supplies (soap, sanitizer, cleaning products)
- Other (anything that doesn't fit above)

UNITS OF MEASURE (use EXACTLY one of these):
- Unit (individual items, single units)
- Case (boxed cases, cartons)
- Box (boxes of items)
- Pack (packs, packages)
- Bag (bags of items)
- lb (pounds - for weight-based items)
- oz (ounces - for smaller weight items)
- kg (kilograms)
- Gallon (gallons of liquid)
- Quart (quarts of liquid)
- Pint (pints)
- Liter (liters)
- Dozen (12 items)
- Bundle (bundles, bunches)
- Roll (rolls of paper, etc.)
- Jar (jars)
- Can (cans)
- Bottle (bottles)

Return the data as valid JSON in this exact format:
{{
    "items": [
        {{
            "item_name": "standardized common name",
            "quantity": number,
            "unit": "one of the units above",
            "unit_price": number or null,
            "supplier_name": "exact supplier name from list or null",
            "category": "one of the categories above"
        }}
    ],
    "order_date": "YYYY-MM-DD or null",
    "order_number": "string or null",
    "total": number or null,
    "supplier_name": "string or null",
    "confidence_score": 0.0 to 1.0
}}

Be thorough and extract ALL items you can identify. If you can't determine a value, use null.
For quantity, always provide a numeric value (default to 1 if unclear).
For unit, ALWAYS pick from the list above - default to "Unit" if unclear.
Always try to assign a category - make your best guess based on the item name.
DO NOT include notes or descriptions for items - just standardize the name itself."""

        if file_type == "docx":
            # For DOCX, extract text and send as text content
            text_content = extract_text_from_docx(content)

            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": f"Please extract all purchase order items from this document text. Return the data as JSON.\n\nDocument content:\n{text_content}"
                    }
                ],
                max_completion_tokens=4096,
                response_format={"type": "json_object"}
            )
        else:
            # For PDF, convert to images first (OpenAI Vision API doesn't support PDFs directly)
            pdf_images = convert_pdf_to_images(content)

            if not pdf_images:
                raise HTTPException(
                    status_code=500,
                    detail="Failed to extract any pages from the PDF"
                )

            # Build the content array with all page images
            content_parts = [
                {
                    "type": "text",
                    "text": f"Please extract all purchase order items from this document ({len(pdf_images)} page(s)). Return the data as JSON."
                }
            ]

            for i, img_base64 in enumerate(pdf_images):
                content_parts.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{img_base64}"
                    }
                })

            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": content_parts
                    }
                ],
                max_completion_tokens=4096,
                response_format={"type": "json_object"}
            )

        # Parse the response
        result_text = response.choices[0].message.content
        result_data = json.loads(result_text)

        # Validate and convert to our schema
        items = []
        for item in result_data.get("items", []):
            items.append(ExtractedOrderItem(
                item_name=item.get("item_name", "Unknown Item"),
                quantity=float(item.get("quantity", 1)),
                unit=item.get("unit"),
                unit_price=float(item["unit_price"]) if item.get("unit_price") else None,
                supplier_name=item.get("supplier_name"),
                category=item.get("category"),
                notes=item.get("notes")
            ))

        return ExtractedOrderData(
            items=items,
            order_date=result_data.get("order_date"),
            order_number=result_data.get("order_number"),
            total=float(result_data["total"]) if result_data.get("total") else None,
            supplier_name=result_data.get("supplier_name"),
            notes=result_data.get("notes"),
            confidence_score=float(result_data.get("confidence_score", 0.8)),
            raw_text=result_text
        )

    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse AI response as JSON: {e}")
        raise HTTPException(
            status_code=500,
            detail="Failed to parse order data from document. The AI response was not valid JSON."
        )
    except Exception as e:
        logger.error(f"Error extracting order from document: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to extract order data from document: {str(e)}"
        )


# ============== ENDPOINTS ==============

@router.post("/extract-order-pdf", response_model=ExtractedOrderData)
async def extract_order_from_document(
    file: UploadFile = File(...),
    current_user: User = Depends(require_admin),
    db: Session = Depends(get_db)
):
    """
    Upload a PDF or DOCX of a historical order and extract item data using AI.
    Returns the extracted data for review before creating the order.
    """
    # Validate file type
    filename_lower = file.filename.lower()
    if filename_lower.endswith('.pdf'):
        file_type = "pdf"
    elif filename_lower.endswith('.docx'):
        file_type = "docx"
    else:
        raise HTTPException(
            status_code=400,
            detail="Only PDF and DOCX files are supported"
        )

    # Read file content
    content = await file.read()

    if len(content) > 10 * 1024 * 1024:  # 10MB limit
        raise HTTPException(
            status_code=400,
            detail="File size exceeds 10MB limit"
        )

    # Fetch existing suppliers to help AI match names
    suppliers = db.query(Supplier).filter(Supplier.is_active == True).all()
    supplier_names = [s.name for s in suppliers]

    # Extract order data using AI
    extracted_data = await extract_order_with_ai(content, file_type, supplier_names)

    return extracted_data


@router.post("/seed-historical-order", response_model=SeedOrderResponse)
def seed_historical_order(
    request: SeedOrderRequest,
    current_user: User = Depends(require_admin),
    db: Session = Depends(get_db)
):
    """
    Create a historical order from extracted/edited data.
    This creates the order with the specified status (default: received).
    """
    # Validate property exists
    property = db.query(Property).filter(Property.id == request.property_id).first()
    if not property:
        raise HTTPException(status_code=404, detail="Property not found")

    # Parse order date
    try:
        order_date = datetime.fromisoformat(request.order_date.replace('Z', '+00:00'))
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format. Use ISO format (YYYY-MM-DD)")

    # Determine status
    status_map = {
        "draft": OrderStatus.DRAFT.value,
        "submitted": OrderStatus.SUBMITTED.value,
        "approved": OrderStatus.APPROVED.value,
        "ordered": OrderStatus.ORDERED.value,
        "received": OrderStatus.RECEIVED.value,
    }
    order_status = status_map.get(request.status.lower(), OrderStatus.RECEIVED.value)

    # Create the order
    order = Order(
        order_number=generate_order_number(property.code),
        property_id=request.property_id,
        week_of=order_date,
        notes=request.notes or f"Historical order seeded on {datetime.utcnow().strftime('%Y-%m-%d')}",
        created_by=current_user.id,
        status=order_status,
        submitted_at=order_date if order_status != OrderStatus.DRAFT.value else None,
        reviewed_by=current_user.id if order_status in [OrderStatus.APPROVED.value, OrderStatus.ORDERED.value, OrderStatus.RECEIVED.value] else None,
        reviewed_at=order_date if order_status in [OrderStatus.APPROVED.value, OrderStatus.ORDERED.value, OrderStatus.RECEIVED.value] else None,
        approved_at=order_date if order_status in [OrderStatus.APPROVED.value, OrderStatus.ORDERED.value, OrderStatus.RECEIVED.value] else None,
        ordered_at=order_date if order_status in [OrderStatus.ORDERED.value, OrderStatus.RECEIVED.value] else None,
        received_at=order_date if order_status == OrderStatus.RECEIVED.value else None,
        created_at=order_date,  # Set created_at to the historical date
    )
    db.add(order)
    db.flush()

    # Add items and create/update inventory items for this property
    for item_data in request.items:
        # Try to find matching supplier from seeded data
        new_supplier_id = None
        if item_data.supplier_name:
            supplier = db.query(Supplier).filter(
                Supplier.name.ilike(f"%{item_data.supplier_name}%")
            ).first()
            if supplier:
                new_supplier_id = supplier.id

        # Use fuzzy matching to find existing inventory item
        inventory_item, match_score = find_matching_inventory_item(
            item_data.item_name,
            request.property_id,
            db,
            similarity_threshold=0.6
        )

        if inventory_item:
            # MATCHED to existing inventory item - use existing data, only fill gaps
            # NEVER overwrite existing data with null values from seeded data

            # Log the match for debugging
            if match_score < 1.0:
                logger.info(f"Fuzzy matched '{item_data.item_name}' -> '{inventory_item.name}' (score: {match_score:.2f})")

            # Only update category if existing is empty AND new data has a value
            if not inventory_item.category and item_data.category:
                inventory_item.category = item_data.category

            # Only update unit_price if existing is empty/zero AND new data has a value
            if (inventory_item.unit_price is None or inventory_item.unit_price == 0) and item_data.unit_price:
                inventory_item.unit_price = item_data.unit_price

            # Only update supplier if existing is empty AND new data has a value
            if not inventory_item.supplier_id and new_supplier_id:
                inventory_item.supplier_id = new_supplier_id

            # Use existing inventory item's values for the order item
            # This ensures consistency with what's already in the system
            final_supplier_id = inventory_item.supplier_id or new_supplier_id
            final_unit = inventory_item.unit or item_data.unit or "Unit"
            final_unit_price = inventory_item.unit_price if inventory_item.unit_price else item_data.unit_price

        else:
            # NO MATCH - create new inventory item
            inventory_item = InventoryItem(
                property_id=request.property_id,
                name=item_data.item_name,
                category=item_data.category,
                supplier_id=new_supplier_id,
                unit=item_data.unit or "Unit",
                unit_price=item_data.unit_price,
                is_active=True,
                is_recurring=item_data.is_recurring
            )
            db.add(inventory_item)
            db.flush()  # Get the ID

            logger.info(f"Created new inventory item: '{item_data.item_name}'")

            final_supplier_id = new_supplier_id
            final_unit = item_data.unit or "Unit"
            final_unit_price = item_data.unit_price

        # Create order item linked to the inventory item
        order_item = OrderItem(
            order_id=order.id,
            inventory_item_id=inventory_item.id,  # Link to matched/created inventory item
            custom_item_name=None,  # Use inventory item's name, not custom
            supplier_id=final_supplier_id,
            flag=OrderItemFlag.MANUAL.value,
            requested_quantity=item_data.quantity,
            approved_quantity=item_data.quantity,  # Auto-approve for historical orders
            received_quantity=item_data.quantity if order_status == OrderStatus.RECEIVED.value else None,
            unit=final_unit,
            unit_price=final_unit_price,
            camp_notes=item_data.notes,
            is_received=order_status == OrderStatus.RECEIVED.value
        )
        db.add(order_item)

    db.commit()
    db.refresh(order)

    # Update estimated total
    order.estimated_total = calculate_order_total(order)
    if order_status == OrderStatus.RECEIVED.value:
        order.actual_total = order.estimated_total
    db.commit()
    db.refresh(order)

    return SeedOrderResponse(
        id=order.id,
        order_number=order.order_number,
        property_id=order.property_id,
        status=order.status,
        week_of=order.week_of,
        item_count=len(order.items),
        estimated_total=order.estimated_total or 0,
        created_at=order.created_at
    )


@router.get("/properties", response_model=List[dict])
def list_properties_for_admin(
    current_user: User = Depends(require_admin),
    db: Session = Depends(get_db)
):
    """List all properties for admin seed dropdown"""
    properties = db.query(Property).filter(Property.is_active == True).all()
    return [{"id": p.id, "name": p.name, "code": p.code} for p in properties]
